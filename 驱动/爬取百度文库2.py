# -*-coding:  utf-8 -*-
# @Time    :  2020/12/22 15:33
# @Author  :  Cooper
# @FileName:  爬取百度文库2.py
# @Software:  PyCharm
from selenium import webdriver
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH# 用来居中显示标题
from time import sleep
from selenium.webdriver.common.keys import Keys

# 浏览器安装路径
#BROWSER_PATH=\'C:\Users\Administrator\AppData\Local\Google\Chrome\Application\chromedriver.exe'
#目的URL
#DEST_URL='https://wenku.baidu.com/view/3bd1e8ef0975f46527d3e14d.html?rec_flag=default'
DEST_URL=input("请输入搜索地址（必须是纯文档）:")
#用来保存文档
doc_title = ''
doc_content_list = []
def find_doc(driver, init=True):
  global doc_content_list
  global doc_title
  global page_count
  stop_condition = False

  if (init is True): # 得到标题
    driver.get(DEST_URL)
    driver.encoding = 'gb2312'
    html = driver.page_source
    soup1 = BeautifulSoup(html, 'lxml')
    page_count = soup1.find(class_ = 'page-count').get_text().split('/')[1]
    title_result = soup1.find('span', attrs={'class': 'doc-header-title'})
    doc_title = title_result.get_text() # 得到文档标题

    sleep(2)
    init = False
    js = "var q=document.documentElement.scrollTop=" + str(3800)
    driver.execute_script(js)
    if soup1.find(class_ = 'div.continue-to-read'):
      element = driver.find_element_by_css_selector('div.continue-to-read')
      webdriver.ActionChains(driver).move_to_element(element).click(element).perform()
      sleep(1)

    for page in range(1,int(page_count)+1):
      input = driver.find_element_by_css_selector('.page-input')
      input.clear()
      input.send_keys(f'{page}')
      input.send_keys(Keys.ENTER)
      sleep(2)
      html = driver.page_source
      soup1 = BeautifulSoup(html, 'html.parser')
      content_result = soup1.find(attrs={'data-page-no': f'{page}'}).get_text()
      content_result = content_result.replace(' ', '')
      doc_content_list.append(content_result)
  print("找不到元素")
  stop_condition = True
          # 得到正文内容
  sleep(2) # 防止页面加载过慢
  if stop_condition is False:
    doc_title, doc_content_list = find_doc(driver, init)
  return doc_title, doc_content_list
def save(doc_title, doc_content_list):
  document = Document()
  heading = document.add_heading(doc_title, 0)
  heading.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中显示
  for each in doc_content_list:
    document.add_paragraph(each)
  # 处理字符编码问题
  t_title = doc_title.split()[0]
  #print(t_title)
  #document.save('4.docx')
  document.save('%s.docx'% t_title)
  print("\n\nCompleted: %s.docx, to read." % t_title)
  driver.quit()
if __name__ == '__main__':
  options = webdriver.ChromeOptions()
  options.add_argument('user-agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.163 Safari/537.36"')
  #driver = webdriver.Chrome(BROWSER_PATH, chrome_options=options)
  driver = webdriver.Chrome()


  #JavascriptExecutor js = (JavascriptExecutor) driver;
  print("**********START**********")
  title, content = find_doc(driver, True)
  save(title, content)
  driver.quit()
